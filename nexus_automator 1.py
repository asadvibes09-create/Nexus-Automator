"""
Nexus Automator - Desktop Automation Tool
A modern desktop automation tool with file organization capabilities
Built with CustomTkinter for a sleek dark-mode UI
"""

import customtkinter as ctk
from tkinter import filedialog, messagebox
import os
import shutil
from pathlib import Path
import threading
from typing import Dict, List
from urllib.parse import urlparse
import requests
from bs4 import BeautifulSoup
import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch

# Set appearance mode and color theme
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")


class NexusAutomator:
    """Main application class for Nexus Automator"""

    FILE_CATEGORIES: Dict[str, List[str]] = {
        "Documents": [".pdf", ".docx", ".doc", ".txt", ".xlsx", ".xls", ".pptx", ".ppt"],
        "Images": [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".svg", ".ico", ".webp"],
        "Videos": [".mp4", ".avi", ".mov", ".mkv", ".flv", ".wmv", ".webm"],
        "Audio": [".mp3", ".wav", ".flac", ".aac", ".ogg", ".wma", ".m4a"],
        "Archives": [".zip", ".rar", ".7z", ".tar", ".gz", ".bz2"],
        "Code": [".py", ".js", ".html", ".css", ".java", ".cpp", ".c", ".go", ".rs"],
        "Executables": [".exe", ".msi", ".bat", ".sh", ".cmd"],
        "Other": []
    }

    def __init__(self, root):
        self.root = root
        self.root.title("Nexus Automator")
        self.root.geometry("900x660")
        self.root.minsize(840, 620)
        self.root.grid_rowconfigure(0, weight=1)
        self.root.grid_columnconfigure(0, weight=1)

        self.selected_folder = None
        self.rename_folder = None
        self.is_organizing = False

        self.create_ui()
        self.select_tab("organizer")

    def create_ui(self):
        outer_frame = ctk.CTkFrame(self.root, fg_color="#0B0F19")
        outer_frame.grid(row=0, column=0, sticky="nsew")
        outer_frame.grid_rowconfigure(0, weight=1)
        outer_frame.grid_columnconfigure(1, weight=1)

        sidebar = ctk.CTkFrame(
            outer_frame,
            fg_color="#070A10",
            width=220,
            corner_radius=0,
            border_width=1,
            border_color="#242C3D"
        )
        sidebar.grid(row=0, column=0, sticky="nswe")
        sidebar.grid_propagate(False)
        sidebar.grid_rowconfigure(7, weight=1)

        logo_label = ctk.CTkLabel(
            sidebar,
            text="Nexus\nAutomator",
            font=("Segoe UI", 20, "bold"),
            text_color="#00E5FF",
            justify="left"
        )
        logo_label.grid(row=0, column=0, padx=24, pady=(26, 16), sticky="w")

        subtitle_label = ctk.CTkLabel(
            sidebar,
            text="Desktop automation made simple.",
            font=("Segoe UI", 11),
            text_color="#7a7a7a",
            justify="left"
        )
        subtitle_label.grid(row=1, column=0, padx=24, sticky="w")

        self.file_tab_btn = ctk.CTkButton(
            sidebar,
            text="📁  File Organizer",
            fg_color="#00E5FF",
            hover_color="#4F46E5",
            text_color="#070A10",
            corner_radius=12,
            font=("Segoe UI", 12, "bold"),
            width=184,
            height=44,
            command=lambda: self.select_tab("organizer")
        )
        self.file_tab_btn.grid(row=2, column=0, padx=18, pady=(20, 8), sticky="ew")

        self.renamer_tab_btn = ctk.CTkButton(
            sidebar,
            text="✏️  Bulk Renamer",
            fg_color="#0B1220",
            hover_color="#1E2638",
            text_color="#94A3B8",
            corner_radius=12,
            font=("Segoe UI", 12, "bold"),
            width=184,
            height=44,
            command=lambda: self.select_tab("renamer")
        )
        self.renamer_tab_btn.grid(row=3, column=0, padx=18, pady=8, sticky="ew")

        self.scraper_tab_btn = ctk.CTkButton(
            sidebar,
            text="🌐  Web Scraper",
            fg_color="#0B1220",
            hover_color="#1E2638",
            text_color="#94A3B8",
            corner_radius=12,
            font=("Segoe UI", 12, "bold"),
            width=184,
            height=44,
            command=lambda: self.select_tab("scraper")
        )
        self.scraper_tab_btn.grid(row=4, column=0, padx=18, pady=8, sticky="ew")

        helper_label = ctk.CTkLabel(
            sidebar,
            text="Select a tool from the sidebar to begin.",
            font=("Segoe UI", 10),
            text_color="#94A3B8",
            wraplength=180,
            justify="left"
        )
        helper_label.grid(row=5, column=0, padx=18, pady=(18, 0), sticky="w")

        content_frame = ctk.CTkFrame(outer_frame, fg_color="#161B26", corner_radius=16)
        content_frame.grid(row=0, column=1, sticky="nsew", padx=(12, 18), pady=18)
        content_frame.grid_rowconfigure(1, weight=1)
        content_frame.grid_columnconfigure(0, weight=1)

        header_card = ctk.CTkFrame(
            content_frame,
            fg_color="#161B26",
            corner_radius=16,
            border_width=1,
            border_color="#242C3D"
        )
        header_card.grid(row=0, column=0, sticky="ew", padx=16, pady=(16, 10))
        header_card.grid_columnconfigure(0, weight=1)

        header_title = ctk.CTkLabel(
            header_card,
            text="Nexus Automator",
            font=("Segoe UI", 24, "bold"),
            text_color="#ffffff",
            anchor="w"
        )
        header_title.grid(row=0, column=0, padx=24, pady=(22, 4), sticky="w")

        header_subtitle = ctk.CTkLabel(
            header_card,
            text="Organize files and rename batches with a polished, dark mode workflow.",
            font=("Segoe UI", 12),
            text_color="#9f9f9f",
            anchor="w"
        )
        header_subtitle.grid(row=1, column=0, padx=24, pady=(0, 22), sticky="w")

        self.content_area = ctk.CTkFrame(content_frame, fg_color="#161B26", corner_radius=16, border_width=1, border_color="#242C3D")
        self.content_area.grid(row=1, column=0, sticky="nsew", padx=16, pady=(0, 10))
        self.content_area.grid_columnconfigure(0, weight=1)

        self.create_organizer_tab()
        self.create_renamer_tab()
        self.create_scraper_tab()

        footer_bar = ctk.CTkFrame(content_frame, fg_color="#1d1d1d", corner_radius=20)
        footer_bar.grid(row=2, column=0, sticky="ew", padx=20, pady=(0, 20))
        footer_bar.grid_columnconfigure(0, weight=1)

        self.progress_bar = ctk.CTkProgressBar(
            footer_bar,
            height=10,
            fg_color="#272727",
            progress_color="#00E5FF"
        )
        self.progress_bar.grid(row=0, column=0, sticky="ew", padx=20, pady=(18, 10))
        self.progress_bar.set(0)

        self.progress_label = ctk.CTkLabel(
            footer_bar,
            text="Ready to organize",
            font=("Segoe UI", 10),
            text_color="#8c8c8c",
            anchor="w"
        )
        self.progress_label.grid(row=1, column=0, sticky="w", padx=20, pady=(0, 18))

    def select_tab(self, tab_name: str):
        self.organizer_frame.grid_remove()
        self.renamer_frame.grid_remove()
        self.scraper_frame.grid_remove()

        inactive_style = {"fg_color": "#0B1220", "text_color": "#94A3B8", "hover_color": "#1E2638"}
        self.file_tab_btn.configure(**inactive_style)
        self.renamer_tab_btn.configure(**inactive_style)
        self.scraper_tab_btn.configure(**inactive_style)

        if tab_name == "organizer":
            self.organizer_frame.grid()
            self.file_tab_btn.configure(fg_color="#00E5FF", text_color="#070A10", hover_color="#4F46E5")
            self.progress_label.configure(text="Ready to organize")
        elif tab_name == "renamer":
            self.renamer_frame.grid()
            self.renamer_tab_btn.configure(fg_color="#00E5FF", text_color="#070A10", hover_color="#4F46E5")
            self.progress_label.configure(text="Ready to rename")
        else:
            self.scraper_frame.grid()
            self.scraper_tab_btn.configure(fg_color="#00E5FF", text_color="#070A10", hover_color="#4F46E5")
            self.progress_label.configure(text="Ready to scrape")

    def create_organizer_tab(self):
        self.organizer_frame = ctk.CTkFrame(self.content_area, fg_color="#1b1b1b", corner_radius=16)
        self.organizer_frame.grid(row=0, column=0, sticky="nsew")
        self.organizer_frame.grid_columnconfigure(0, weight=1)

        title = ctk.CTkLabel(
            self.organizer_frame,
            text="File Organizer",
            font=("Segoe UI", 18, "bold"),
            text_color="#ffffff",
            anchor="w"
        )
        title.grid(row=0, column=0, sticky="w", padx=24, pady=(22, 6))

        subtitle = ctk.CTkLabel(
            self.organizer_frame,
            text="Select a folder, then organize files into extension-based categories.",
            font=("Segoe UI", 11),
            text_color="#a9a9a9",
            anchor="w"
        )
        subtitle.grid(row=1, column=0, sticky="w", padx=24)

        folder_card = ctk.CTkFrame(self.organizer_frame, fg_color="#1E2638", corner_radius=12, border_width=1, border_color="#242C3D")
        folder_card.grid(row=2, column=0, sticky="ew", padx=18, pady=(18, 14))
        folder_card.grid_columnconfigure(0, weight=1)
        folder_card.grid_columnconfigure(1, weight=0)

        self.folder_path_label = ctk.CTkLabel(
            folder_card,
            text="No folder selected",
            font=("Segoe UI", 11),
            text_color="#9f9f9f",
            anchor="w"
        )
        self.folder_path_label.grid(row=0, column=0, sticky="ew", padx=18, pady=18)

        browse_btn = ctk.CTkButton(
            folder_card,
            text="Browse",
            width=120,
            corner_radius=12,
            fg_color="#00d4ff",
            hover_color="#0099d6",
            text_color="#000000",
            font=("Segoe UI", 11, "bold"),
            command=self.browse_folder
        )
        browse_btn.grid(row=0, column=1, sticky="e", padx=18, pady=18)

        details_card = ctk.CTkFrame(self.organizer_frame, fg_color="#1E2638", corner_radius=12, border_width=1, border_color="#242C3D")
        details_card.grid(row=3, column=0, sticky="ew", padx=18, pady=(0, 14))
        details_card.grid_columnconfigure(0, weight=1)

        details_title = ctk.CTkLabel(
            details_card,
            text="Organization Details",
            font=("Segoe UI", 13, "bold"),
            text_color="#ffffff",
            anchor="w"
        )
        details_title.grid(row=0, column=0, sticky="w", padx=18, pady=(18, 4))

        categories_text = "Files will be organized into:\n"
        for category, extensions in self.FILE_CATEGORIES.items():
            if extensions:
                ext_list = ", ".join(extensions[:3])
                if len(extensions) > 3:
                    ext_list += f", +{len(extensions)-3} more"
                categories_text += f"• {category}: {ext_list}\n"
        categories_text += "• Other: Unknown file types"

        self.info_label = ctk.CTkLabel(
            details_card,
            text=categories_text,
            font=("Segoe UI", 10),
            text_color="#c4c4c4",
            justify="left"
        )
        self.info_label.grid(row=1, column=0, sticky="ew", padx=18, pady=(0, 18))

        actions_card = ctk.CTkFrame(self.organizer_frame, fg_color="#1E2638", corner_radius=12, border_width=1, border_color="#242C3D")
        actions_card.grid(row=4, column=0, sticky="ew", padx=18, pady=(0, 14))
        actions_card.grid_columnconfigure(0, weight=1)
        actions_card.grid_columnconfigure(1, weight=1)

        self.organize_btn = ctk.CTkButton(
            actions_card,
            text="🚀 Organize Files",
            fg_color="#00d4ff",
            hover_color="#0099d6",
            text_color="#000000",
            font=("Segoe UI", 12, "bold"),
            corner_radius=12,
            command=self.organize_files,
            height=46
        )
        self.organize_btn.grid(row=0, column=0, sticky="ew", padx=(18, 8), pady=18)

        reset_btn = ctk.CTkButton(
            actions_card,
            text="🔄 Reset",
            fg_color="#2a2a2a",
            hover_color="#333333",
            text_color="#ffffff",
            font=("Segoe UI", 12, "bold"),
            corner_radius=12,
            command=self.reset_ui,
            height=46
        )
        reset_btn.grid(row=0, column=1, sticky="ew", padx=(8, 18), pady=18)

        self.organizer_status = ctk.CTkLabel(
            self.organizer_frame,
            text="",
            font=("Segoe UI", 10),
            text_color="#00d4ff",
            anchor="w"
        )
        self.organizer_status.grid(row=5, column=0, sticky="w", padx=24, pady=(0, 22))

    def create_renamer_tab(self):
        self.renamer_frame = ctk.CTkFrame(self.content_area, fg_color="#161B26", corner_radius=16)
        self.renamer_frame.grid(row=0, column=0, sticky="nsew")
        self.renamer_frame.grid_columnconfigure(0, weight=1)

        title = ctk.CTkLabel(
            self.renamer_frame,
            text="Bulk Renamer",
            font=("Segoe UI", 18, "bold"),
            text_color="#ffffff",
            anchor="w"
        )
        title.grid(row=0, column=0, sticky="w", padx=24, pady=(22, 6))

        subtitle = ctk.CTkLabel(
            self.renamer_frame,
            text="Choose a folder, type a new prefix, then rename files in sequence.",
            font=("Segoe UI", 11),
            text_color="#a9a9a9",
            anchor="w"
        )
        subtitle.grid(row=1, column=0, sticky="w", padx=24)

        rename_card = ctk.CTkFrame(self.renamer_frame, fg_color="#1E2638", corner_radius=12, border_width=1, border_color="#242C3D")
        rename_card.grid(row=2, column=0, sticky="ew", padx=18, pady=(18, 14))
        rename_card.grid_columnconfigure(0, weight=1)
        rename_card.grid_columnconfigure(1, weight=0)

        self.rename_folder_label = ctk.CTkLabel(
            rename_card,
            text="No folder selected",
            font=("Segoe UI", 11),
            text_color="#9f9f9f",
            anchor="w"
        )
        self.rename_folder_label.grid(row=0, column=0, sticky="ew", padx=18, pady=18)

        browse_btn = ctk.CTkButton(
            rename_card,
            text="Browse",
            width=120,
            corner_radius=12,
            fg_color="#00d4ff",
            hover_color="#0099d6",
            text_color="#000000",
            font=("Segoe UI", 11, "bold"),
            command=self.browse_renamer_folder
        )
        browse_btn.grid(row=0, column=1, sticky="e", padx=18, pady=18)

        prefix_card = ctk.CTkFrame(self.renamer_frame, fg_color="#1E2638", corner_radius=12, border_width=1, border_color="#242C3D")
        prefix_card.grid(row=3, column=0, sticky="ew", padx=18, pady=(0, 14))
        prefix_card.grid_columnconfigure(0, weight=1)

        prefix_label = ctk.CTkLabel(
            prefix_card,
            text="New Prefix Name",
            font=("Segoe UI", 11, "bold"),
            text_color="#ffffff",
            anchor="w"
        )
        prefix_label.grid(row=0, column=0, sticky="w", padx=18, pady=(18, 6))

        self.prefix_entry = ctk.CTkEntry(
            prefix_card,
            placeholder_text="e.g. ProjectX",
            font=("Segoe UI", 11),
            fg_color="#181818",
            text_color="#ffffff",
            width=380,
            height=42,
            corner_radius=12
        )
        self.prefix_entry.grid(row=1, column=0, sticky="ew", padx=18, pady=(0, 18))

        execute_btn = ctk.CTkButton(
            self.renamer_frame,
            text="Execute Rename",
            fg_color="#00d4ff",
            hover_color="#0099d6",
            text_color="#000000",
            font=("Segoe UI", 12, "bold"),
            corner_radius=12,
            command=self.execute_rename,
            height=46
        )
        execute_btn.grid(row=4, column=0, sticky="w", padx=24, pady=(10, 12))

        self.renamer_status = ctk.CTkLabel(
            self.renamer_frame,
            text="",
            font=("Segoe UI", 10),
            text_color="#00E5FF",
            anchor="w"
        )
        self.renamer_status.grid(row=5, column=0, sticky="w", padx=24, pady=(0, 22))

    def create_scraper_tab(self):
        self.scraper_frame = ctk.CTkFrame(self.content_area, fg_color="#161B26", corner_radius=16)
        self.scraper_frame.grid(row=0, column=0, sticky="nsew")
        self.scraper_frame.grid_columnconfigure(0, weight=1)

        title = ctk.CTkLabel(
            self.scraper_frame,
            text="Web Scraper",
            font=("Segoe UI", 18, "bold"),
            text_color="#ffffff",
            anchor="w"
        )
        title.grid(row=0, column=0, sticky="w", padx=24, pady=(22, 6))

        subtitle = ctk.CTkLabel(
            self.scraper_frame,
            text="Extract headings and links from a webpage and save to Excel, CSV, Word, or PDF.",
            font=("Segoe UI", 11),
            text_color="#a9a9a9",
            anchor="w"
        )
        subtitle.grid(row=1, column=0, sticky="w", padx=24)

        scraper_card = ctk.CTkFrame(self.scraper_frame, fg_color="#1E2638", corner_radius=12, border_width=1, border_color="#242C3D")
        scraper_card.grid(row=2, column=0, sticky="ew", padx=18, pady=(18, 14))
        scraper_card.grid_columnconfigure(0, weight=1)
        scraper_card.grid_columnconfigure(1, weight=0)

        self.url_entry = ctk.CTkEntry(
            scraper_card,
            placeholder_text="Enter website URL (e.g., https://example.com)",
            font=("Segoe UI", 11),
            fg_color="#1E2638",
            border_color="#242C3D",
            text_color="#ffffff",
            width=540,
            height=44,
            corner_radius=12
        )
        self.url_entry.grid(row=0, column=0, sticky="ew", padx=18, pady=18, columnspan=2)

        format_card = ctk.CTkFrame(self.scraper_frame, fg_color="#1E2638", corner_radius=12, border_width=1, border_color="#242C3D")
        format_card.grid(row=3, column=0, sticky="ew", padx=18, pady=(0, 14))
        format_card.grid_columnconfigure(0, weight=1)

        format_label = ctk.CTkLabel(
            format_card,
            text="Export Format",
            font=("Segoe UI", 11, "bold"),
            text_color="#ffffff",
            anchor="w"
        )
        format_label.grid(row=0, column=0, sticky="w", padx=18, pady=(18, 6))

        self.format_option = ctk.CTkOptionMenu(
            format_card,
            values=["Excel (.xlsx)", "CSV (.csv)", "Word (.docx)", "PDF (.pdf)"],
            fg_color="#1E2638",
            button_color="#242C3D",
            dropdown_hover_color="#1E2638",
            text_color="#ffffff",
            button_hover_color="#4F46E5",
            font=("Segoe UI", 11),
            width=200,
            corner_radius=12
        )
        self.format_option.set("Excel (.xlsx)")
        self.format_option.grid(row=1, column=0, sticky="w", padx=18, pady=(0, 18))

        self.extract_btn = ctk.CTkButton(
            self.scraper_frame,
            text="Extract Data",
            fg_color="#00E5FF",
            hover_color="#4F46E5",
            text_color="#070A10",
            font=("Segoe UI", 12, "bold"),
            corner_radius=12,
            command=self.execute_scrape,
            height=46
        )
        self.extract_btn.grid(row=4, column=0, sticky="w", padx=24, pady=(10, 12))

        self.scraper_status = ctk.CTkLabel(
            self.scraper_frame,
            text="",
            font=("Segoe UI", 10),
            text_color="#00d4ff",
            anchor="w"
        )
        self.scraper_status.grid(row=5, column=0, sticky="w", padx=24, pady=(0, 22))

    def browse_folder(self):
        folder = filedialog.askdirectory(title="Select a folder to organize")
        if folder:
            self.selected_folder = folder
            display_path = folder if len(folder) < 40 else "..." + folder[-37:]
            self.folder_path_label.configure(text=display_path, text_color="#00d4ff")
            self.organizer_status.configure(text="✓ Folder selected. Ready to organize!")
            self.progress_label.configure(text="Ready to organize")

    def browse_renamer_folder(self):
        folder = filedialog.askdirectory(title="Select a folder for bulk renaming")
        if folder:
            self.rename_folder = folder
            display_path = folder if len(folder) < 40 else "..." + folder[-37:]
            self.rename_folder_label.configure(text=display_path, text_color="#00d4ff")
            self.renamer_status.configure(text="✓ Folder selected. Ready to rename!")
            self.progress_label.configure(text="Ready to rename")

    def organize_files(self):
        if not self.selected_folder:
            messagebox.showwarning("Warning", "Please select a folder first!")
            return

        if not os.path.exists(self.selected_folder):
            messagebox.showerror("Error", "Selected folder no longer exists!")
            return

        self.is_organizing = True
        self.organize_btn.configure(state="disabled", fg_color="#666666")
        threading.Thread(target=self._organize_thread, daemon=True).start()

    def execute_rename(self):
        prefix = self.prefix_entry.get().strip()
        if not self.rename_folder:
            messagebox.showwarning("Warning", "Please select a folder for renaming first.")
            return

        if not prefix:
            messagebox.showwarning("Warning", "Please enter a new prefix name.")
            return

        if not os.path.exists(self.rename_folder):
            messagebox.showerror("Error", "Selected folder no longer exists!")
            return

        files = [f for f in os.listdir(self.rename_folder) if os.path.isfile(os.path.join(self.rename_folder, f))]
        if not files:
            messagebox.showinfo("Info", "No files found to rename in this folder.")
            return

        renamed_count = 0
        for index, filename in enumerate(files, start=1):
            source_path = os.path.join(self.rename_folder, filename)
            name, extension = os.path.splitext(filename)
            new_name = f"{prefix}_{index}{extension}"
            dest_path = os.path.join(self.rename_folder, new_name)

            counter = 1
            while os.path.exists(dest_path):
                dest_path = os.path.join(self.rename_folder, f"{prefix}_{index}_{counter}{extension}")
                counter += 1

            try:
                os.rename(source_path, dest_path)
                renamed_count += 1
            except Exception as e:
                print(f"Failed to rename {filename}: {e}")

        self.renamer_status.configure(text=f"✓ Renamed {renamed_count} files successfully.")
        messagebox.showinfo("Rename Complete", f"Successfully renamed {renamed_count} files.")
        self.progress_label.configure(text="Bulk rename completed")

    def execute_scrape(self):
        url = self.url_entry.get().strip()
        if not url:
            messagebox.showwarning("Warning", "Please enter a website URL to scrape.")
            return

        if not url.startswith("http"):
            url = "https://" + url

        self.scraper_status.configure(text="✓ Fetching page... this may take a moment.")
        self.progress_bar.set(0.1)
        self.extract_btn.configure(state="disabled", fg_color="#666666")
        self.root.after(0, lambda: self.progress_label.configure(text="Scraping webpage..."))

        chosen_format = self.format_option.get()
        threading.Thread(target=self._scrape_thread, args=(url, chosen_format), daemon=True).start()

    def _scrape_thread(self, url: str, chosen_format: str):
        try:
            response = requests.get(url, timeout=15, headers={"User-Agent": "Mozilla/5.0"})
            response.raise_for_status()
            soup = BeautifulSoup(response.text, "html.parser")

            headings = [
                {"type": tag.name, "text": tag.get_text(strip=True)}
                for tag in soup.find_all(["h1", "h2", "h3"])
            ]

            links = [
                {"text": tag.get_text(strip=True), "url": tag["href"]}
                for tag in soup.find_all("a", href=True)
            ]

            safe_name = self._safe_website_name(url)
            desktop_path = Path.home() / "Desktop"
            if not desktop_path.exists():
                desktop_path = Path.cwd()

            if chosen_format.startswith("Excel"):
                rows = []
                max_rows = max(len(headings), len(links), 1)
                for index in range(max_rows):
                    row = {
                        "Heading Type": "",
                        "Heading Text": "",
                        "Link Text": "",
                        "Link URL": ""
                    }
                    if index < len(headings):
                        row["Heading Type"] = headings[index]["type"]
                        row["Heading Text"] = headings[index]["text"]
                    if index < len(links):
                        row["Link Text"] = links[index]["text"]
                        row["Link URL"] = links[index]["url"]
                    rows.append(row)

                data_frame = pd.DataFrame(rows)
                output_file = desktop_path / f"scraped_data_{safe_name}.xlsx"
                data_frame.to_excel(output_file, index=False)
            elif chosen_format.startswith("CSV"):
                rows = []
                max_rows = max(len(headings), len(links), 1)
                for index in range(max_rows):
                    row = {
                        "Heading Type": "",
                        "Heading Text": "",
                        "Link Text": "",
                        "Link URL": ""
                    }
                    if index < len(headings):
                        row["Heading Type"] = headings[index]["type"]
                        row["Heading Text"] = headings[index]["text"]
                    if index < len(links):
                        row["Link Text"] = links[index]["text"]
                        row["Link URL"] = links[index]["url"]
                    rows.append(row)

                data_frame = pd.DataFrame(rows)
                output_file = desktop_path / f"scraped_data_{safe_name}.csv"
                data_frame.to_csv(output_file, index=False)
            elif chosen_format.startswith("Word"):
                output_file = desktop_path / f"scraped_data_{safe_name}.docx"
                self._save_to_word(output_file, url, headings, links)
            else:
                output_file = desktop_path / f"scraped_data_{safe_name}.pdf"
                self._save_to_pdf(output_file, url, headings, links)

            self.root.after(0, lambda: self.scraper_complete(output_file))
        except Exception as exc:
            self.root.after(0, lambda: self.scraper_error(str(exc)))

    def scraper_complete(self, output_file: Path):
        self.scraper_status.configure(text=f"✓ Data saved to {output_file.name}")
        self.extract_btn.configure(state="normal", fg_color="#00d4ff")
        messagebox.showinfo("Scrape Complete", f"Web data exported successfully to {output_file}")
        self.progress_label.configure(text="Web scraping completed")
        self.progress_bar.set(1.0)

    def scraper_error(self, error_message: str):
        self.scraper_status.configure(text="✗ Scraping failed. Check your URL and connection.")
        self.extract_btn.configure(state="normal", fg_color="#00d4ff")
        messagebox.showerror("Scrape Error", f"Failed to scrape webpage:\n{error_message}")
        self.progress_label.configure(text="Scrape failed")

    def _safe_website_name(self, url: str) -> str:
        parsed = urlparse(url)
        domain = parsed.netloc or parsed.path
        safe_name = "".join(c if c.isalnum() else "_" for c in domain)
        return safe_name or "website"

    def _save_to_word(self, output_file: Path, url: str, headings: list, links: list):
        document = Document()
        document.add_heading(f"Scraped Data from {url}", level=1)
        document.add_paragraph("\nWebsite Headings", style="IntenseQuote")

        if headings:
            for heading in headings:
                p = document.add_paragraph(style="List Bullet")
                run = p.add_run(f"{heading['type'].upper()}: ")
                run.bold = True
                p.add_run(heading["text"])
        else:
            document.add_paragraph("No headings found.")

        document.add_paragraph("\nExtracted Links", style="IntenseQuote")
        if links:
            for link in links:
                p = document.add_paragraph(style="List Number")
                run = p.add_run(link["text"] or link["url"])
                run.bold = True
                p.add_run(f" — {link['url']}")
        else:
            document.add_paragraph("No hyperlinks found.")

        document.save(output_file)

    def _save_to_pdf(self, output_file: Path, url: str, headings: list, links: list):
        pdf_canvas = canvas.Canvas(str(output_file), pagesize=letter)
        width, height = letter
        text_object = pdf_canvas.beginText(40, height - 40)
        text_object.setFont("Helvetica-Bold", 16)
        text_object.textLine(f"Scraped Data from {url}")
        text_object.moveCursor(0, 14)
        text_object.setFont("Helvetica", 11)
        text_object.textLine("")
        text_object.setFont("Helvetica-Bold", 12)
        text_object.textLine("Website Headings")
        text_object.setFont("Helvetica", 10)

        if headings:
            for heading in headings:
                heading_text = f"{heading['type'].upper()}: {heading['text']}"
                self._pdf_add_wrapped_text(text_object, pdf_canvas, heading_text, width - 80)
        else:
            text_object.textLine("No headings found.")

        text_object.textLine("")
        text_object.setFont("Helvetica-Bold", 12)
        text_object.textLine("Extracted Links")
        text_object.setFont("Helvetica", 10)

        if links:
            for link in links:
                link_text = f"{link['text'] or link['url']} — {link['url']}"
                self._pdf_add_wrapped_text(text_object, pdf_canvas, link_text, width - 80)
        else:
            text_object.textLine("No hyperlinks found.")

        pdf_canvas.drawText(text_object)
        pdf_canvas.save()

    def _pdf_add_wrapped_text(self, text_object, pdf_canvas, content: str, max_width: float):
        words = content.split()
        line = ""
        for word in words:
            test_line = f"{line} {word}".strip()
            if pdf_canvas.stringWidth(test_line, "Helvetica", 10) > max_width:
                if line:
                    text_object.textLine(line)
                line = word
            else:
                line = test_line
        if line:
            text_object.textLine(line)

    def _organize_thread(self):
        try:
            files = os.listdir(self.selected_folder)
            total_files = len([f for f in files if os.path.isfile(os.path.join(self.selected_folder, f))])

            if total_files == 0:
                self.root.after(0, lambda: messagebox.showinfo("Info", "No files to organize in this folder!"))
                self.root.after(0, self.reset_ui)
                return

            organized_count = 0
            for filename in files:
                file_path = os.path.join(self.selected_folder, filename)
                if os.path.isdir(file_path):
                    continue

                file_ext = Path(filename).suffix.lower()
                category = "Other"
                for cat, extensions in self.FILE_CATEGORIES.items():
                    if file_ext in extensions:
                        category = cat
                        break

                category_folder = os.path.join(self.selected_folder, category)
                os.makedirs(category_folder, exist_ok=True)

                destination = os.path.join(category_folder, filename)
                if os.path.exists(destination):
                    name, ext = os.path.splitext(filename)
                    counter = 1
                    while os.path.exists(os.path.join(category_folder, f"{name}_{counter}{ext}")):
                        counter += 1
                    destination = os.path.join(category_folder, f"{name}_{counter}{ext}")

                try:
                    shutil.move(file_path, destination)
                    organized_count += 1
                except Exception as e:
                    print(f"Error moving {filename}: {str(e)}")

                progress = organized_count / total_files if total_files else 0
                self.root.after(0, lambda p=progress, o=organized_count, t=total_files: self.update_progress(p, o, t))

            self.root.after(0, lambda: self.show_success(organized_count, total_files))
        except Exception as e:
            self.root.after(0, lambda: messagebox.showerror("Error", f"An error occurred: {str(e)}"))
        finally:
            self.root.after(0, self.reset_ui)
            self.is_organizing = False

    def update_progress(self, progress, organized, total):
        self.progress_bar.set(progress)
        self.progress_label.configure(text=f"Organizing... {organized}/{total} files")

    def show_success(self, organized_count, total_files):
        message = f"✓ Success!\n\n{organized_count} out of {total_files} files organized successfully!"
        messagebox.showinfo("Organization Complete", message)
        self.organizer_status.configure(text=f"✓ Successfully organized {organized_count} files!")
        self.progress_label.configure(text="Organization complete")

    def reset_ui(self):
        self.organize_btn.configure(state="normal", fg_color="#00d4ff")
        self.progress_bar.set(0)
        self.progress_label.configure(text="Ready to organize")


def main():
    root = ctk.CTk()
    app = NexusAutomator(root)
    root.mainloop()


if __name__ == "__main__":
    main()
